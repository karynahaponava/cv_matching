import os
import re
import io
import docx
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

load_dotenv()

SCOPES = [
    "https://www.googleapis.com/auth/documents.readonly",
    "https://www.googleapis.com/auth/drive.readonly",
]

_DOC_ID_PATTERNS = (
    re.compile(r"/document/d/([a-zA-Z0-9_-]+)"),
    re.compile(r"/file/d/([a-zA-Z0-9_-]+)"),
    re.compile(r"[?&]id=([a-zA-Z0-9_-]+)"),
)


def extract_doc_id(url: str) -> str | None:
    if not url:
        return None
    for pattern in _DOC_ID_PATTERNS:
        match = pattern.search(url)
        if match:
            return match.group(1)
    return None


def _get_credentials():
    creds_path = os.getenv("GOOGLE_CREDS_PATH")
    if not creds_path:
        raise ValueError("GOOGLE_CREDS_PATH is not set")
    return service_account.Credentials.from_service_account_file(
        creds_path, scopes=SCOPES
    )


def _read_paragraph_element(element: dict) -> str:
    text_run = element.get("textRun")
    if not text_run:
        return ""
    return text_run.get("content", "")


def _read_structural_elements(elements: list) -> str:
    text = ""
    for value in elements:
        if "paragraph" in value:
            for elem in value["paragraph"].get("elements", []):
                text += _read_paragraph_element(elem)
        elif "table" in value:
            for row in value["table"].get("tableRows", []):
                for cell in row.get("tableCells", []):
                    text += _read_structural_elements(cell.get("content", []))
        elif "tableOfContents" in value:
            text += _read_structural_elements(
                value["tableOfContents"].get("content", [])
            )
    return text


def _fetch_via_docs_api(doc_id: str, credentials) -> str:
    service = build("docs", "v1", credentials=credentials, cache_discovery=False)
    document = service.documents().get(documentId=doc_id).execute()
    content = document.get("body", {}).get("content", [])
    return _read_structural_elements(content).strip()


def _fetch_via_drive_export(doc_id: str, credentials) -> str:
    service = build("drive", "v3", credentials=credentials, cache_discovery=False)
    meta = service.files().get(fileId=doc_id, fields="mimeType").execute()
    mime_type = meta.get("mimeType", "")

    if mime_type == "application/vnd.google-apps.document":
        raw = service.files().export(fileId=doc_id, mimeType="text/plain").execute()
        if isinstance(raw, bytes):
            return raw.decode("utf-8").strip()
        return str(raw).strip()

    elif mime_type == "application/pdf":
        request = service.files().get_media(fileId=doc_id)
        file_content = request.execute()
        reader = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()

    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        request = service.files().get_media(fileId=doc_id)
        file_content = request.execute()
        doc = docx.Document(io.BytesIO(file_content))
        
        full_text = [p.text for p in doc.paragraphs]
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
                        
        return "\n".join(full_text)

    return ""


def get_doc_text(doc_url: str) -> str:
    doc_id = extract_doc_id(doc_url)
    if not doc_id:
        return ""

    credentials = _get_credentials()

    try:
        text = _fetch_via_drive_export(doc_id, credentials)
        if text:
            return text
    except Exception as e:
        print(f"⚠️ Ошибка Drive API для {doc_id}: {e}")

    try:
        return _fetch_via_docs_api(doc_id, credentials)
    except Exception as e:
        return ""
